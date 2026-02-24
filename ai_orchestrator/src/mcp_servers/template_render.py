"""MCP tool server: DOCX, PDF, PPTX, and CSV document rendering."""
import csv
import io
import logging
import os
from typing import Any

logger = logging.getLogger("ai_orchestrator.mcp.template_render")

_MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "http://minio:9000")
_MINIO_ACCESS = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
_MINIO_SECRET = os.getenv("MINIO_SECRET_KEY", "minioadmin")
_BUCKET = os.getenv("MINIO_BUCKET", "deal-manager")


async def render_proposal_section_docx(
    section_title: str,
    content: str,
    template_name: str = "standard_proposal",
    diagrams: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Render a proposal section as a DOCX file.

    Args:
        section_title: Section title (e.g. "Technical Approach").
        content: Markdown or plain text content.
        template_name: Template to use ("standard_proposal", "volume_1", "volume_2").
        diagrams: Optional list of diagram dicts with image_data (base64) and caption.

    Returns:
        Dict with object_key (MinIO), url, filename, size_bytes.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
        import base64

        doc = Document()

        # Heading
        heading = doc.add_heading(section_title, level=1)

        # Content – simple paragraph-based rendering from markdown
        for paragraph in content.split("\n\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            if paragraph.startswith("## "):
                doc.add_heading(paragraph[3:], level=2)
            elif paragraph.startswith("### "):
                doc.add_heading(paragraph[4:], level=3)
            elif paragraph.startswith("- ") or paragraph.startswith("* "):
                for line in paragraph.split("\n"):
                    line = line.lstrip("-* ").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            elif paragraph.startswith("1. ") or paragraph.startswith("1) "):
                for i, line in enumerate(paragraph.split("\n"), 1):
                    line = line.lstrip("0123456789.) ").strip()
                    if line:
                        doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(paragraph)

        # Embed diagrams
        if diagrams:
            for diag in diagrams:
                img_b64 = diag.get("image_data", "")
                caption = diag.get("caption", "")
                if img_b64:
                    img_bytes = base64.b64decode(img_b64)
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                    if caption:
                        doc.add_paragraph(caption, style="Caption")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        filename = f"{section_title.replace(' ', '_')[:50]}.docx"
        upload_result = await _upload_bytes(docx_bytes, f"proposals/{filename}", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        return {**upload_result, "filename": filename}

    except ImportError as exc:
        logger.warning("python-docx not installed: %s", exc)
        return {
            "content": content,
            "format": "text",
            "note": "python-docx not available; returning plain text",
        }
    except Exception as exc:
        logger.error("DOCX render failed: %s", exc)
        return {"error": str(exc)}


async def render_full_proposal_docx(
    proposal_id: str,
    sections: list[dict[str, Any]],
    cover_page: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Render a complete multi-section proposal as a single DOCX file.

    Args:
        proposal_id: Proposal UUID for naming.
        sections: List of section dicts with title, content, level, and optional diagrams.
        cover_page: Optional cover page dict with opportunity_name, company_name, date.

    Returns:
        Dict with object_key, url, filename, section_count, page_estimate.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        import base64

        doc = Document()

        # Cover page
        if cover_page:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cover_page.get("opportunity_name", "Proposal"))
            run.bold = True
            run.font.size = Pt(24)
            doc.add_paragraph()
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(f"Prepared by: {cover_page.get('company_name', '')}")
            p2.add_run(f"\nDate: {cover_page.get('date', '')}")
            doc.add_page_break()

        # Table of contents placeholder
        doc.add_heading("Table of Contents", level=1)
        for sec in sections:
            doc.add_paragraph(sec.get("title", ""), style="List Bullet")
        doc.add_page_break()

        # Sections
        for sec in sections:
            level = sec.get("level", 1)
            doc.add_heading(sec.get("title", "Section"), level=level)
            content = sec.get("content", "")
            for para in content.split("\n\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

            for diag in sec.get("diagrams", []):
                img_b64 = diag.get("image_data", "")
                if img_b64:
                    img_bytes = base64.b64decode(img_b64)
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
                    if diag.get("caption"):
                        doc.add_paragraph(diag["caption"], style="Caption")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        filename = f"proposal_{proposal_id[:8]}.docx"
        upload_result = await _upload_bytes(
            docx_bytes, f"proposals/{filename}",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return {
            **upload_result,
            "filename": filename,
            "section_count": len(sections),
            "page_estimate": max(1, len(sections) * 3),
        }

    except Exception as exc:
        logger.error("Full proposal DOCX render failed: %s", exc)
        return {"error": str(exc), "proposal_id": proposal_id}


async def render_csv(
    data: list[dict[str, Any]],
    filename: str = "export.csv",
) -> dict[str, Any]:
    """Render a list of dicts as a CSV file.

    Args:
        data: List of row dicts (all rows must share the same keys).
        filename: Output filename.

    Returns:
        Dict with object_key, url, filename, row_count.
    """
    if not data:
        return {"error": "No data provided", "row_count": 0}

    try:
        buf = io.StringIO()
        writer = csv.DictWriter(buf, fieldnames=list(data[0].keys()))
        writer.writeheader()
        writer.writerows(data)
        csv_bytes = buf.getvalue().encode("utf-8")

        upload_result = await _upload_bytes(csv_bytes, f"exports/{filename}", "text/csv")
        return {**upload_result, "filename": filename, "row_count": len(data)}
    except Exception as exc:
        logger.error("CSV render failed: %s", exc)
        return {"error": str(exc)}


async def render_pricing_volume(
    deal_id: str,
    pricing_scenario: dict[str, Any],
    rate_schedule: list[dict[str, Any]],
    loe_estimate: dict[str, Any],
) -> dict[str, Any]:
    """Render a complete Cost/Pricing Volume (Volume IV) as DOCX.

    Args:
        deal_id: Deal UUID.
        pricing_scenario: Selected pricing scenario dict.
        rate_schedule: Labor category rate schedule list.
        loe_estimate: LOE estimate by task/labor category.

    Returns:
        Dict with object_key, url, filename.
    """
    sections = [
        {
            "title": "1. Price Summary",
            "level": 1,
            "content": _format_price_summary(pricing_scenario),
            "diagrams": [],
        },
        {
            "title": "2. Basis of Estimate (BOE)",
            "level": 1,
            "content": _format_boe(loe_estimate),
            "diagrams": [],
        },
        {
            "title": "3. Labor Rate Schedule",
            "level": 1,
            "content": _format_rate_table(rate_schedule),
            "diagrams": [],
        },
        {
            "title": "4. Staffing Plan",
            "level": 1,
            "content": _format_staffing(loe_estimate),
            "diagrams": [],
        },
    ]
    return await render_full_proposal_docx(
        proposal_id=deal_id,
        sections=sections,
        cover_page={"opportunity_name": "Cost Volume", "company_name": ""},
    )


# ── Internal helpers ──────────────────────────────────────────────────────────

async def _upload_bytes(file_bytes: bytes, object_key: str, content_type: str) -> dict:
    try:
        from src.mcp_servers.document_tools import upload_to_storage

        return await upload_to_storage(file_bytes, object_key, content_type)
    except Exception as exc:
        logger.warning("Upload failed, returning in-memory result: %s", exc)
        import base64

        return {
            "object_key": object_key,
            "url": None,
            "size_bytes": len(file_bytes),
            "data_b64": base64.b64encode(file_bytes).decode(),
        }


def _format_price_summary(scenario: dict) -> str:
    lines = [
        f"**Total Price: ${scenario.get('total_price', 0):,.2f}**",
        f"Strategy: {scenario.get('strategy_type', 'competitive')}",
        f"Estimated Win Probability: {scenario.get('win_probability', 0)*100:.0f}%",
        f"Expected Value: ${scenario.get('expected_value', 0):,.2f}",
        f"Profit Margin: {scenario.get('margin_percent', 0):.1f}%",
    ]
    return "\n\n".join(lines)


def _format_boe(loe: dict) -> str:
    tasks = loe.get("tasks", [])
    if not tasks:
        return "Basis of estimate details not available."
    lines = []
    for task in tasks:
        lines.append(f"**{task.get('name', 'Task')}**: {task.get('hours', 0)} hours")
        lines.append(f"  Labor categories: {', '.join(task.get('labor_categories', []))}")
    return "\n\n".join(lines)


def _format_rate_table(rates: list[dict]) -> str:
    if not rates:
        return "Labor rate schedule not provided."
    lines = ["| Labor Category | Level | Fully Loaded Rate/Hour |", "|---|---|---|"]
    for r in rates:
        lines.append(
            f"| {r.get('category', '')} | {r.get('level', '')} | ${r.get('fully_loaded_rate', 0):,.2f} |"
        )
    return "\n".join(lines)


def _format_staffing(loe: dict) -> str:
    total_hours = loe.get("total_hours", 0)
    total_fte = loe.get("total_fte", 0)
    period = loe.get("period_of_performance", "")
    return (
        f"Total Labor Hours: {total_hours:,}\n\n"
        f"Average FTE: {total_fte:.1f}\n\n"
        f"Period of Performance: {period}"
    )
