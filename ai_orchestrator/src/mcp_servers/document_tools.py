"""MCP tool server: PDF/DOCX parsing, chunking, and embedding pipeline."""
import io
import logging
import os
from pathlib import Path
from typing import Any

logger = logging.getLogger("ai_orchestrator.mcp.documents")

_MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "http://minio:9000")
_MINIO_ACCESS = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
_MINIO_SECRET = os.getenv("MINIO_SECRET_KEY", "minioadmin")
_BUCKET = os.getenv("MINIO_BUCKET", "deal-manager")


# ── Parsing helpers ───────────────────────────────────────────────────────────

def _parse_pdf(data: bytes) -> str:
    """Extract raw text from a PDF byte string."""
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        logger.warning("pdfplumber not installed; falling back to PyPDF2")
    try:
        from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(
            page.extract_text() or "" for page in reader.pages
        )
    except Exception as exc:
        logger.error("PDF parse failed: %s", exc)
        return ""


def _parse_docx(data: bytes) -> str:
    """Extract raw text from a DOCX byte string."""
    try:
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also include table text
        for table in doc.tables:
            for row in table.rows:
                cells = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if cells:
                    paragraphs.append(cells)
        return "\n\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX parse failed: %s", exc)
        return ""


def _parse_excel(data: bytes) -> list[dict]:
    """Parse an Excel file and return rows as a list of dicts."""
    try:
        import openpyxl  # type: ignore

        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        rows_all: list[dict] = []
        for sheet in wb.worksheets:
            headers = [str(cell.value or "") for cell in next(sheet.iter_rows(max_row=1))]
            for row in sheet.iter_rows(min_row=2, values_only=True):
                rows_all.append(dict(zip(headers, [str(v or "") for v in row])))
        return rows_all
    except Exception as exc:
        logger.error("Excel parse failed: %s", exc)
        return []


# ── Public MCP tools ──────────────────────────────────────────────────────────

async def parse_document(
    file_path: str | None = None,
    file_bytes: bytes | None = None,
    filename: str = "",
) -> dict[str, Any]:
    """Parse a document file (PDF, DOCX, or Excel) and return its text content.

    Provide either *file_path* (MinIO object key or local path) or *file_bytes*.

    Returns:
        Dict with keys: text (str), page_count (int), filename (str).
    """
    if file_bytes is None:
        if file_path is None:
            return {"error": "Must provide file_path or file_bytes"}
        file_bytes = await _load_from_minio(file_path)
        filename = filename or Path(file_path).name

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = _parse_pdf(file_bytes)
        page_count = max(1, text.count("\n\n"))
    elif ext in (".docx", ".doc"):
        text = _parse_docx(file_bytes)
        page_count = max(1, len(text) // 3000)
    elif ext in (".xlsx", ".xls"):
        rows = _parse_excel(file_bytes)
        text = "\n".join(" | ".join(r.values()) for r in rows)
        page_count = len(rows)
    else:
        text = file_bytes.decode("utf-8", errors="replace")
        page_count = max(1, len(text) // 3000)

    return {"text": text, "page_count": page_count, "filename": filename, "char_count": len(text)}


async def chunk_and_embed(
    text: str,
    source_id: str = "",
    content_type: str = "text",
) -> list[dict[str, Any]]:
    """Chunk *text* and generate embeddings for each chunk.

    Returns:
        List of dicts with keys: text, chunk_index, source_id, embedding (list[float]).
    """
    from src.rag.chunker import chunk_document
    from src.rag.embeddings import embed_batch

    chunks = chunk_document(text, source_id=source_id, content_type=content_type)  # type: ignore
    if not chunks:
        return []

    texts = [c.text for c in chunks]
    embeddings = await embed_batch(texts)

    result = []
    for chunk, embedding in zip(chunks, embeddings):
        result.append(
            {
                "text": chunk.text,
                "chunk_index": chunk.chunk_index,
                "source_id": chunk.source_id,
                "content_type": chunk.content_type,
                "metadata": chunk.metadata,
                "embedding": embedding,
            }
        )
    return result


async def parse_and_chunk(
    file_path: str | None = None,
    file_bytes: bytes | None = None,
    filename: str = "",
    source_id: str = "",
) -> dict[str, Any]:
    """Full pipeline: parse a document and return chunks with embeddings.

    Returns:
        Dict with keys: filename, text, chunks (list), chunk_count (int).
    """
    parsed = await parse_document(
        file_path=file_path, file_bytes=file_bytes, filename=filename
    )
    if "error" in parsed:
        return parsed

    chunks = await chunk_and_embed(
        parsed["text"],
        source_id=source_id or parsed.get("filename", ""),
    )
    return {
        "filename": parsed["filename"],
        "text": parsed["text"],
        "chunks": chunks,
        "chunk_count": len(chunks),
        "char_count": parsed["char_count"],
    }


async def upload_to_storage(
    file_bytes: bytes,
    object_key: str,
    content_type: str = "application/octet-stream",
) -> dict[str, Any]:
    """Upload *file_bytes* to MinIO object storage.

    Returns:
        Dict with keys: object_key, url, size_bytes.
    """
    try:
        from minio import Minio  # type: ignore

        client = Minio(
            _MINIO_ENDPOINT.replace("http://", "").replace("https://", ""),
            access_key=_MINIO_ACCESS,
            secret_key=_MINIO_SECRET,
            secure=_MINIO_ENDPOINT.startswith("https"),
        )
        if not client.bucket_exists(_BUCKET):
            client.make_bucket(_BUCKET)

        client.put_object(
            _BUCKET,
            object_key,
            io.BytesIO(file_bytes),
            length=len(file_bytes),
            content_type=content_type,
        )
        return {
            "object_key": object_key,
            "url": f"{_MINIO_ENDPOINT}/{_BUCKET}/{object_key}",
            "size_bytes": len(file_bytes),
        }
    except Exception as exc:
        logger.error("MinIO upload failed for %s: %s", object_key, exc)
        return {"error": str(exc), "object_key": object_key}


# ── Internal helpers ──────────────────────────────────────────────────────────

async def _load_from_minio(object_key: str) -> bytes:
    """Download *object_key* from MinIO and return raw bytes."""
    try:
        from minio import Minio  # type: ignore

        client = Minio(
            _MINIO_ENDPOINT.replace("http://", "").replace("https://", ""),
            access_key=_MINIO_ACCESS,
            secret_key=_MINIO_SECRET,
            secure=_MINIO_ENDPOINT.startswith("https"),
        )
        response = client.get_object(_BUCKET, object_key)
        return response.read()
    except Exception as exc:
        logger.error("MinIO download failed for %s: %s", object_key, exc)
        # Try local filesystem as fallback
        p = Path(object_key)
        if p.exists():
            return p.read_bytes()
        return b""
